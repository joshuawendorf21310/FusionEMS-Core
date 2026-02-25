from __future__ import annotations

from pathlib import Path
from datetime import datetime

def create_docx(out_dir: Path, title: str, body: str) -> Path:
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)
    name = f"{_safe(title)}-{_ts()}.docx"
    path = out_dir / name
    doc.save(str(path))
    return path

def create_xlsx(out_dir: Path, title: str, body: str) -> Path:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = title
    ws["A2"] = body
    name = f"{_safe(title)}-{_ts()}.xlsx"
    path = out_dir / name
    wb.save(str(path))
    return path

def create_invoice_pdf(out_dir: Path, title: str, body: str) -> Path:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    name = f"{_safe(title)}-{_ts()}.pdf"
    path = out_dir / name
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, height - 72, title)
    c.setFont("Helvetica", 11)
    y = height - 110
    for line in body.splitlines()[:80]:
        c.drawString(72, y, line[:110])
        y -= 14
        if y < 72:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 72
    c.showPage()
    c.save()
    return path

def _ts() -> str:
    return datetime.utcnow().strftime("%Y%m%d%H%M%S")

def _safe(s: str) -> str:
    return "".join(ch for ch in s.strip().replace(" ", "_") if ch.isalnum() or ch in ("_", "-"))[:40] or "document"
